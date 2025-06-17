from django.http import HttpResponse
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from urllib.parse import quote
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime


def get_kz_month_name(month_number):
    kz_month_names = {
        1: 'Қаңтар',
        2: 'Ақпан',
        3: 'Наурыз',
        4: 'Сәуір',
        5: 'Мамыр',
        6: 'Маусым',
        7: 'Шілде',
        8: 'Тамыз',
        9: 'Қыркүйек',
        10: 'Қазан',
        11: 'Қараша',
        12: 'Желтоқсан',
    }
    return kz_month_names.get(month_number, "")

def create_first_page(doc, faculty_name, year):
    """Функция для создания первой страницы отчета"""
    next_year = year + 1  # Следующий учебный год

    # --- Первая страница ---
    doc.add_paragraph("Қожа Ахмет Ясауи атындағы Халықаралық қазақ-түрік университеті").alignment = WD_ALIGN_PARAGRAPH.CENTER

    para = doc.add_paragraph(
        "«БЕКІТЕМІН»\n"
        "Сапа бойынша басшылық өкілі, Ғылым және стратегиялық даму вице-ректоры\n"
        "__________________________ А.Ошибаева\n"
        f"«____» _______________ {year}ж."
    )
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Пустые строки
    for _ in range(6):
        doc.add_paragraph("")

    doc.add_paragraph(f"{faculty_name}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"{year} - {next_year} оқу жылына").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("ИНДИКАТИВТІ ЖОСПАРЫ").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\nКентау").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Добавляем разрыв страницы для следующей части отчета
    doc.add_page_break()



def init_document(selected_year, faculty):
    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.text = f"Жүктеу күні мен уақыты: {datetime.now().strftime('%d-%m-%Y %H:%M:%S')}"

    document.add_paragraph("Қожа Ахмет Ясауи атындағы Халықаралық қазақ-түрік университеті").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(
        "«БЕКІТЕМІН»\nСапа бойынша басшылық өкілі, Ғылым және стратегиялық даму вице-ректоры\n"
        f"__________________________ А.Ошибаева\n«____» _______________ {selected_year.year}ж."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for _ in range(6):
        document.add_paragraph("")

    document.add_paragraph(f"{faculty.name}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"{selected_year.year} - {selected_year.year + 1} оқу жылына").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("ИНДИКАТИВТІ ЖОСПАРЫ").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("\nТүркістан").alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    return document

def add_direction_title(document, selected_year, faculty, direction):
    for _ in range(6):
        document.add_paragraph("")
    document.add_paragraph(f"{faculty.name}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"{selected_year.year} - {selected_year.year + 1} оқу жылына").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("ИНДИКАТИВТІ ЖОСПАРЫ").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(f"{direction.id}  {direction.name.upper()}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()

def create_indicator_table(document, departments):
    # Создаем таблицу с 3 + количеством департаментов столбцов
    table = document.add_table(rows=1, cols=3 + len(departments))
    table.style = 'Table Grid'

    cols = table.columns

    # Устанавливаем ширину столбцов
    cols[0].width = Inches(1)  # Код (меньше)
    cols[1].width = Inches(5)  # Индикатор атауы (больше)
    for idx in range(2, len(cols) - 1):  # столбцы для департаментов
        cols[idx].width = Inches(2)
    cols[-1].width = Inches(1)  # Сумма (меньше)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Код'
    hdr_cells[1].text = 'Индикатор атауы'
    for idx, dept in enumerate(departments):
        hdr_cells[2 + idx].text = dept.name
    hdr_cells[-1].text = 'Сумма'

    # Выравнивание текста в заголовках по центру
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return table

def set_cell_font(cell, bold=False, align_center=False):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            run.font.size = Pt(12)
            run.font.bold = bold
            run.font.color.rgb = None  # чёрный

def add_footer(section):
    """
    Время в нижнем колонтитуле
    """
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    current_time = datetime.now().strftime('%d-%m-%Y %H:%M:%S')
    paragraph.text = f"Жүктеу күні мен уақыты: {current_time}"

def generate_docx_response(doc, teacher, year):
    """
    Генерирует HttpResponse с Word-документом для скачивания.

    :param doc: Объект Document (python-docx)
    :param teacher: Объект учителя
    :param year: Объект года
    :return: HttpResponse с документом
    """
    file_name = f"Мұғалім {teacher.last_name} {teacher.first_name} - {year.year}ж.docx"
    encoded_file_name = quote(file_name)

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename*=UTF-8\'\'{encoded_file_name}'

    doc.save(response)
    return response

def configure_page(doc):
    """
    Настраивает ориентацию и отступы страницы документа Word.

    :param doc: объект docx.Document
    """
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # Меняем местами ширину и высоту, т.к. ландшафт
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

def align_cell_center(cell):
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# -----====== Dean ======--------
def generate_dean(document, filename_base):
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    filename = f"{filename_base}.docx"
    encoded_file_name = quote(filename)
    response['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_file_name}"
    document.save(response)
    return response
