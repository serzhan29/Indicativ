from django.db.models import Sum
from django.http import HttpResponse
from docx import Document
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404
from docx.shared import Pt
from .models import AggregatedIndicator, Direction, Year, Indicator, TeacherReport, User, MainIndicator
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from django.views import View
from user.models import Department, Faculty
from docx.shared import Cm
from django.db import models
from .docx_utils import (
    create_first_page, get_kz_month_name, init_document, add_direction_title,
    create_indicator_table, set_cell_font, add_footer, generate_docx_response,
    configure_page, align_cell_center, generate_dean
)


@login_required
def download_teacher_report(request, teacher_id, direction_id, year_id):
    """Генерация отчета в Word с правильным форматированием"""
    teacher = get_object_or_404(User, id=teacher_id)
    direction = get_object_or_404(Direction, id=direction_id)
    year = get_object_or_404(Year, id=year_id)
    next_year = year.year + 1  # Следующий учебный год
    faculty_name = teacher.profile.faculty.name if teacher.profile.faculty else "Көрсетілмеген"

    aggregated_data = AggregatedIndicator.objects.filter(
        teacher=teacher, year=year, main_indicator__direction=direction
    )

    doc = Document()
    # --- Первая страница ---
    create_first_page(doc, faculty_name, year.year)

    for _ in range(6):
        doc.add_paragraph("")

    doc.add_paragraph(f"{faculty_name}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"{year.year} - {next_year} оқу жылына").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("ИНДИКАТИВТІ ЖОСПАРЫ").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"{direction.id}  {direction.name.upper()}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Страница с таблицей ---
    doc.add_page_break()
    # --- Настройка страницы ---
    configure_page(doc)
    # Время
    section = doc.sections[-1]
    add_footer(section)

    # --- Таблица ---
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.autofit = False  # отключаем автоширину

    # Установка ширины столбцов
    column_widths = [Cm(1.5), Cm(15), Cm(2), Cm(2), Cm(2.5), Cm(2)]
    for i, width in enumerate(column_widths):
        table.columns[i].width = width
        table.rows[0].cells[i].width = width

    # Заголовки
    hdr_cells = table.rows[0].cells
    headers = ["Код", "Индикатор атауы", "Өлшем бірлігі", "Сумма", "Дедлайн (ай/жыл)", "Құжат"]
    for idx, text in enumerate(headers):
        hdr_cells[idx].text = text
        hdr_cells[idx].width = column_widths[idx]
        hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hdr_cells[idx].paragraphs[0].runs[0]
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    # Данные
    for data in aggregated_data:
        # Дедлайн для агрегированных данных
        deadline_main = "-"
        if data.deadline_month and data.deadline_year:
            month_name = get_kz_month_name(data.deadline_month)
            deadline_main = f"{month_name} / {data.deadline_year}"

        # Проверяем наличие загруженных документов
        has_docs = "✓" if hasattr(data, 'uploaded_works') and data.uploaded_works.exists() else "-"

        row = table.add_row().cells
        row_data = [
            data.main_indicator.code or "-",
            data.main_indicator.name,
            data.main_indicator.unit,
            str(data.total_value),
            deadline_main,
            has_docs,
        ]
        for idx, value in enumerate(row_data):
            row[idx].text = value
            row[idx].width = column_widths[idx]
            run = row[idx].paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            row[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
            run.bold = True

            # Подиндикаторы
        sub_indicators = Indicator.objects.filter(main_indicator=data.main_indicator, years=year)
        for ind in sub_indicators:
            # Получаем отчет по подиндикатору
            report = TeacherReport.objects.filter(teacher=teacher, indicator=ind, year=year).first()

            # Дедлайн для подиндикатора
            deadline_sub = "-"
            if report and report.deadline_month and report.deadline_year:
                month_name = get_kz_month_name(report.deadline_month)
                deadline_sub = f"{month_name} / {report.deadline_year}"

            has_docs_sub = "✓" if report and report.uploaded_works.exists() else "-"

            sub_row = table.add_row().cells
            sub_row_data = [
                ind.code or "-",
                ind.name,
                ind.unit,
                str(report.value) if report else "0",
                deadline_sub,
                has_docs_sub,
            ]
            for idx, value in enumerate(sub_row_data):
                sub_row[idx].text = value
                sub_row[idx].width = column_widths[idx]
                run = sub_row[idx].paragraphs[0].runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                sub_row[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 1 else WD_ALIGN_PARAGRAPH.CENTER

    # Возвращаем файл
    return generate_docx_response(doc, teacher, year)


class TeacherReportWordExportView(View):
    """ Для всех направлений """
    def get(self, request, *args, **kwargs):
        year_id = request.GET.get('year')
        teacher_id = request.GET.get('teacher')

        if not year_id:
            return HttpResponse("Оқу жылы көрсетілмеген", status=400)

        year = get_object_or_404(Year, id=year_id)
        teacher = get_object_or_404(User, id=teacher_id) if teacher_id else request.user
        directions = Direction.objects.all().order_by('id')
        next_year = year.year + 1

        faculty_name = (teacher.profile.faculty.name if hasattr(teacher, 'profile') and teacher.profile.faculty else "Факультет")

        doc = Document()
        configure_page(doc)
        create_first_page(doc, faculty_name, year.year)
        section = doc.sections[-1]
        add_footer(section)

        for direction in directions:
            for _ in range(6):
                doc.add_paragraph("")

            doc.add_paragraph(f"{faculty_name}").alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"{year.year} - {next_year} оқу жылына").alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("ИНДИКАТИВТІ ЖОСПАРЫ").alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"{direction.id}  {direction.name.upper()}").alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_page_break()
            main_indicators = MainIndicator.objects.filter(direction=direction, years=year)

            if not main_indicators.exists():
                doc.add_paragraph("Индикаторлар бойынша деректер жоқ.")
                continue

            # Создание таблицы с фиксированной шириной столбцов
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            table.autofit = False  # отключаем автоширину

            # Устанавливаем ширины колонок
            column_widths = [Cm(1.5), Cm(15), Cm(2), Cm(2), Cm(2.5), Cm(2)]
            for i, width in enumerate(column_widths):
                table.columns[i].width = width
                table.rows[0].cells[i].width = width

            # Заголовки
            hdr_cells = table.rows[0].cells
            headers = ['Код', 'Индикатор атауы', 'Өлшем бірлігі', 'Мәні', 'Дедлайн (ай/жыл)', 'Құжат']

            for i, header_text in enumerate(headers):
                hdr_cells[i].text = header_text
                set_cell_font(hdr_cells[i], bold=True, align_center=True)


            for main_indicator in main_indicators:
                aggregated_data = AggregatedIndicator.objects.filter(
                    teacher=teacher, main_indicator=main_indicator, year=year
                ).first()

                total_value = aggregated_data.total_value if aggregated_data else 0
                total_value = "0" if total_value == 0 else str(total_value)

                deadline_main = f"{get_kz_month_name(aggregated_data.deadline_month)} / {aggregated_data.deadline_year}" if aggregated_data and aggregated_data.deadline_month and aggregated_data.deadline_year else "-"
                confirm_symbol = "✓" if aggregated_data and aggregated_data.uploaded_works.exists() else "-"


                row = table.add_row().cells
                values = [
                    main_indicator.code,
                    main_indicator.name,
                    main_indicator.unit,
                    str(total_value),
                    deadline_main,
                    confirm_symbol,
                ]

                align_centers = [True, False, True, True, True, True]

                for i, (value, align_center) in enumerate(zip(values, align_centers)):
                    row[i].text = value
                    row[i].width = column_widths[i]
                    set_cell_font(row[i], bold=True, align_center=align_center)

                indicators = Indicator.objects.filter(main_indicator=main_indicator, years=year)

                for indicator in indicators:
                    report = TeacherReport.objects.filter(
                        teacher=teacher, indicator=indicator, year=year
                    ).first()

                    value = report.value if report else 0
                    deadline_sub = f"{get_kz_month_name(report.deadline_month)} / {report.deadline_year}" if report and report.deadline_month and report.deadline_year else "-"
                    confirm_symbol = "✓" if report and report.uploaded_works.exists() else "-"

                    ind_row = table.add_row().cells
                    values = [
                        f'{indicator.code}',
                        indicator.name,
                        indicator.unit,
                        str(value),
                        deadline_sub,
                        confirm_symbol,
                    ]

                    align_centers = [True, False, True, True, True, True]

                    for i, (value, align_center) in enumerate(zip(values, align_centers)):
                        ind_row[i].text = value
                        ind_row[i].width = column_widths[i]
                        set_cell_font(ind_row[i], align_center=align_center)

            if direction != directions.last():
                doc.add_page_break()

        return generate_docx_response(doc, teacher, year)


def export_report(request, faculty_id):
    """ с учителем """
    year_id = request.GET.get('year')
    department_id = request.GET.get('department')  # Получаем id кафедры из GET параметров

    if year_id:
        selected_year = Year.objects.get(id=year_id)
    else:
        selected_year = Year.objects.latest('year')

    faculty = Faculty.objects.get(id=faculty_id)

    # Если передан department_id, берем только одну кафедру, иначе все кафедры факультета
    if department_id:
        departments = Department.objects.filter(id=department_id)
    else:
        departments = Department.objects.filter(faculty=faculty)

    directions = Direction.objects.all()

    document = init_document(selected_year, faculty)

    for direction in directions:
        add_direction_title(document, selected_year, faculty, direction)

        table = create_indicator_table(document, departments)

        main_indicators = MainIndicator.objects.filter(direction=direction, years=selected_year).prefetch_related("indicators")

        for main in main_indicators:
            has_sub_indicators = main.indicators.exists()

            if has_sub_indicators:
                sub_indicators = main.indicators.filter(years=selected_year)
                dept_sums = [0] * len(departments)

                for sub in sub_indicators:
                    for idx, dept in enumerate(departments):
                        value = TeacherReport.objects.filter(
                            indicator=sub,
                            year=selected_year,
                            teacher__profile__department=dept,
                            value__gte=1
                        ).aggregate(total=models.Sum('value'))['total'] or 0
                        dept_sums[idx] += value

                summary_row = table.add_row().cells
                summary_row[0].text = f'{main.code}'
                summary_row[1].text = f'{main.name}'
                summary_row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                summary_row[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for idx, value in enumerate(dept_sums):
                    summary_row[2 + idx].text = str(value)
                summary_row[-1].text = str(sum(dept_sums))


                summary_row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                summary_row[-1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                for sub in sub_indicators:
                    row = table.add_row().cells
                    row[0].text = sub.code
                    row[1].text = sub.name

                    sub_total = 0
                    for idx, dept in enumerate(departments):
                        reports = TeacherReport.objects.filter(
                            indicator=sub,
                            year=selected_year,
                            teacher__profile__department=dept,
                            value__gte=1
                        )
                        cell_text = "\n".join(f"{r.teacher.get_full_name()}: {r.value}" for r in reports)
                        value_sum = sum(r.value for r in reports)
                        sub_total += value_sum
                        row[2 + idx].text = cell_text if cell_text else "—"

                    row[-1].text = str(sub_total)

                    # Выравнивание для "Код" и "Сумма"
                    row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    row[-1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            else:
                row = table.add_row().cells
                row[0].text = main.code
                row[1].text = main.name

                total = 0
                for idx, dept in enumerate(departments):
                    aggr_reports = AggregatedIndicator.objects.filter(
                        main_indicator=main,
                        year=selected_year,
                        teacher__profile__department=dept,
                        total_value__gte=1
                    )
                    values = [(r.teacher.get_full_name(), r.total_value) for r in aggr_reports]
                    cell_text = "\n".join(f"{t[0]}: {t[1]}" for t in values)
                    row[2 + idx].text = cell_text if cell_text else "—"
                    total += sum(val for _, val in values)

                row[-1].text = str(total)

                # Выравнивание для "Код" и "Сумма"
                row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                row[-1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if direction != directions.last():
            document.add_page_break()

    filename_base = f"Факультет есебі - {faculty.name} {selected_year.year} (мұғаліммен)"
    return generate_dean(document, filename_base)


@login_required
def export_department_report_docx(request, faculty_id):
    year_id = request.GET.get('year')
    selected_year = Year.objects.get(id=year_id) if year_id else Year.objects.latest('year')

    faculty = Faculty.objects.get(id=faculty_id)

    department_id = request.GET.get('department')
    if department_id and department_id != 'all':
        departments = Department.objects.filter(id=department_id, faculty=faculty)
    else:
        departments = Department.objects.filter(faculty=faculty)

    directions = Direction.objects.all()
    document = init_document(selected_year, faculty)

    for direction in directions:
        add_direction_title(document, selected_year, faculty, direction)

        num_depts = len(departments)
        table = document.add_table(rows=1, cols=3 + num_depts)
        table.style = 'Table Grid'

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Код'
        hdr_cells[1].text = 'Индикатор атауы'
        for idx, dept in enumerate(departments):
            hdr_cells[2 + idx].text = dept.name
        hdr_cells[-1].text = 'Жалпы сумма'

        column_widths = [Cm(1), Cm(23)] + [Cm(1) for _ in departments] + [Cm(1)]
        for idx, cell in enumerate(hdr_cells):
            cell.width = column_widths[idx]
            align_cell_center(cell)

        main_indicators = MainIndicator.objects.filter(
            direction=direction, years=selected_year
        ).prefetch_related("indicators")

        for main in main_indicators:
            has_sub_indicators = main.indicators.exists()

            if has_sub_indicators:
                sub_indicators = main.indicators.filter(years=selected_year)
                dept_sums = [0] * num_depts

                for sub in sub_indicators:
                    for idx, dept in enumerate(departments):
                        value = TeacherReport.objects.filter(
                            indicator=sub,
                            year=selected_year,
                            teacher__profile__department=dept
                        ).aggregate(total=Sum('value'))['total'] or 0
                        dept_sums[idx] += value

                summary_row = table.add_row().cells

                # Код (жирный)
                p_code = summary_row[0].paragraphs[0]
                run_code = p_code.add_run(f'{main.code}')
                run_code.bold = True
                align_cell_center(summary_row[0])

                # Название (жирный)
                p = summary_row[1].paragraphs[0]
                run = p.add_run(f'Барлығы: {main.name}')
                run.bold = True

                for idx, value in enumerate(dept_sums):
                    summary_row[2 + idx].text = str(value)
                summary_row[-1].text = str(sum(dept_sums))

                for idx in range(2, len(summary_row)):
                    align_cell_center(summary_row[idx])

                for sub in sub_indicators:
                    row = table.add_row().cells
                    row[0].text = sub.code
                    row[1].text = sub.name

                    sub_total = 0
                    for idx, dept in enumerate(departments):
                        value = TeacherReport.objects.filter(
                            indicator=sub,
                            year=selected_year,
                            teacher__profile__department=dept
                        ).aggregate(total=Sum('value'))['total'] or 0
                        row[2 + idx].text = str(value)
                        sub_total += value

                    row[-1].text = str(sub_total)

                    align_cell_center(row[0])
                    for idx in range(2, len(row)):
                        align_cell_center(row[idx])

            else:
                row = table.add_row().cells

                # Код (жирный)
                p_code = row[0].paragraphs[0]
                run_code = p_code.add_run(main.code)
                run_code.bold = True
                align_cell_center(row[0])

                # Название (жирный)
                p = row[1].paragraphs[0]
                run = p.add_run(main.name)
                run.bold = True

                total = 0
                for idx, dept in enumerate(departments):
                    value = AggregatedIndicator.objects.filter(
                        main_indicator=main,
                        year=selected_year,
                        teacher__profile__department=dept
                    ).aggregate(total=Sum('total_value'))['total'] or 0
                    row[2 + idx].text = str(value)
                    total += value

                row[-1].text = str(total)

                for idx in range(2, len(row)):
                    align_cell_center(row[idx])

        if direction != directions.last():
            document.add_page_break()

    filename_base = f"Факультет есебі - {faculty.name} {selected_year.year} (мұғалімсіз)"
    return generate_dean(document, filename_base)


